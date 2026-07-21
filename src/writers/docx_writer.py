"""DOCX translate-in-place while keeping paragraph/run formatting."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

ProgressCb = Callable[[int, int, str], None]


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for p in section.footer.paragraphs:
            yield p
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace paragraph text, keeping the first run's formatting."""
    runs = paragraph.runs
    if not runs:
        if new_text:
            paragraph.add_run(new_text)
        return

    # Keep formatting of the first run that had content (or first run)
    target_idx = 0
    for i, run in enumerate(runs):
        if run.text:
            target_idx = i
            break

    runs[target_idx].text = new_text
    for i, run in enumerate(runs):
        if i != target_idx:
            run.text = ""


def collect_docx_texts(path: Path | str) -> list[str]:
    doc = Document(str(path))
    texts: list[str] = []
    for p in _iter_paragraphs(doc):
        t = _paragraph_text(p).strip()
        if t:
            texts.append(t)
    return texts


def translate_docx(
    input_path: Path | str,
    output_path: Path | str,
    translate_fn: Callable[[list[str]], list[str]],
    progress: ProgressCb | None = None,
) -> Path:
    """
    Translate DOCX in place.

    `translate_fn` receives a list of non-empty paragraph strings and returns
    translations in the same order.
    """
    input_path = Path(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(input_path))
    targets: list[Paragraph] = []
    sources: list[str] = []
    for p in _iter_paragraphs(doc):
        t = _paragraph_text(p)
        if t.strip():
            targets.append(p)
            sources.append(t)

    total = len(sources)
    if progress:
        progress(0, max(total, 1), "DOCX matnlar tarjima qilinmoqda…")

    if sources:
        # Chunk to avoid huge batches
        translated: list[str] = []
        batch_size = 16
        for start in range(0, len(sources), batch_size):
            chunk = sources[start : start + batch_size]
            translated.extend(translate_fn(chunk))
            if progress:
                progress(min(start + len(chunk), total), total, "DOCX tarjima…")
    else:
        translated = []

    for paragraph, new_text in zip(targets, translated):
        _set_paragraph_text(paragraph, new_text)

    doc.save(str(output_path))
    if progress:
        progress(total, max(total, 1), "DOCX saqlandi")
    return output_path
