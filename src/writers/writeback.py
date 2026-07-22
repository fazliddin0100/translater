"""Write translated units back to DOCX or PDF (images preserved in PDF)."""

from __future__ import annotations

import io
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.utils.normalize import normalize_uzbek_text
from src.utils.units import TextUnit

_FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\arial.ttf"),
    Path(r"C:\Windows\Fonts\segoeui.ttf"),
    Path(r"C:\Windows\Fonts\calibri.ttf"),
]


def _pick_font() -> str | None:
    for p in _FONT_CANDIDATES:
        if p.exists():
            return str(p)
    return None


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for p in section.footer.paragraphs:
            yield p
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        if new_text:
            paragraph.add_run(new_text)
        return
    target_idx = 0
    for i, run in enumerate(runs):
        if run.text:
            target_idx = i
            break
    runs[target_idx].text = new_text
    for i, run in enumerate(runs):
        if i != target_idx:
            run.text = ""


def write_docx(input_path: Path | str, output_path: Path | str, units: list[TextUnit]) -> Path:
    """Replace DOCX texts in place (images/styles in original DOCX stay)."""
    input_path = Path(input_path)
    output_path = Path(output_path).with_suffix(".docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    by_index = {u.meta["index"]: u for u in units if u.kind == "docx_para"}
    doc = Document(str(input_path))

    idx = 0
    for paragraph in _iter_paragraphs(doc):
        text = _paragraph_text(paragraph)
        if not text.strip():
            continue
        unit = by_index.get(idx)
        if unit is not None:
            _set_paragraph_text(
                paragraph, normalize_uzbek_text(unit.translated or unit.text)
            )
        idx += 1

    doc.save(str(output_path))
    return output_path


def _page_images(pdf: fitz.Document, page_index: int) -> list[tuple[float, bytes]]:
    """Return (y0, png_bytes) for images on a page, top-to-bottom."""
    page = pdf[page_index]
    out: list[tuple[float, bytes]] = []
    seen: set[int] = set()
    try:
        infos = page.get_image_info(xrefs=True)
    except Exception:
        infos = []

    for info in infos:
        xref = info.get("xref")
        if not xref or xref in seen:
            continue
        seen.add(xref)
        bbox = info.get("bbox") or (0, 0, 0, 0)
        y0 = float(bbox[1])
        try:
            pix = fitz.Pixmap(pdf, xref)
            if pix.n - pix.alpha >= 4:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            if pix.width < 8 or pix.height < 8:
                continue
            out.append((y0, pix.tobytes("png")))
        except Exception:
            continue
    out.sort(key=lambda t: t[0])
    return out


def write_units_as_new_docx(
    output_path: Path | str,
    units: list[TextUnit],
    *,
    source_pdf: Path | str | None = None,
) -> Path:
    """
    Build editable Word from translated units.
    If source_pdf is given, also copy page images into the Word file.
    """
    output_path = Path(output_path).with_suffix(".docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    by_page: dict[int, list[TextUnit]] = {}
    for u in units:
        page = int(u.meta.get("page", 0)) if u.kind == "pdf_span" else 0
        by_page.setdefault(page, []).append(u)

    pdf = fitz.open(str(source_pdf)) if source_pdf else None
    try:
        pages = sorted(by_page.keys())
        if pdf is not None:
            pages = list(range(len(pdf)))

        for i, page_i in enumerate(pages):
            if i > 0:
                doc.add_page_break()
            doc.add_heading(f"Sahifa {page_i + 1}", level=2)

            items: list[tuple[float, str, object]] = []
            for u in by_page.get(page_i, []):
                bbox = u.meta.get("bbox") or [0, 0, 0, 0]
                y = float(bbox[1]) if len(bbox) > 1 else 0.0
                text = normalize_uzbek_text(u.translated or u.text or "").strip()
                if text:
                    items.append((y, "text", text))

            if pdf is not None:
                for y, png in _page_images(pdf, page_i):
                    items.append((y, "image", png))

            items.sort(key=lambda t: t[0])
            for _y, kind, payload in items:
                if kind == "text":
                    doc.add_paragraph(str(payload))
                else:
                    stream = io.BytesIO(payload)  # type: ignore[arg-type]
                    try:
                        doc.add_picture(stream, width=Inches(5.8))
                    except Exception:
                        pass
    finally:
        if pdf is not None:
            pdf.close()

    doc.save(str(output_path))
    return output_path


def _fit_fontsize(rect: fitz.Rect, text: str, start: float) -> float:
    size = max(4.0, float(start))
    while size >= 4.0:
        avg_char = size * 0.5
        lines = max(1, int(rect.height / (size * 1.2)))
        max_chars = max(1, int((rect.width * lines) / avg_char))
        if len(text) <= max_chars * 1.2:
            return size
        size -= 0.5
    return 4.0


def write_pdf(input_path: Path | str, output_path: Path | str, units: list[TextUnit]) -> Path:
    """
    Replace text spans in PDF; images and drawings stay (only text boxes are redacted).
    """
    input_path = Path(input_path)
    output_path = Path(output_path).with_suffix(".pdf")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    fontfile = _pick_font()
    doc = fitz.open(str(input_path))
    pdf_units = [u for u in units if u.kind == "pdf_span"]

    by_page: dict[int, list[TextUnit]] = {}
    for u in pdf_units:
        by_page.setdefault(int(u.meta["page"]), []).append(u)

    for page_index, page_units in by_page.items():
        page = doc[page_index]
        for u in page_units:
            page.add_redact_annot(fitz.Rect(u.meta["bbox"]), fill=(1, 1, 1))
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        for u in page_units:
            new_text = normalize_uzbek_text(u.translated or u.text).strip()
            if not new_text:
                continue
            rect = fitz.Rect(u.meta["bbox"])
            size = float(u.meta.get("size") or 11)
            rect.y1 = max(rect.y1, rect.y0 + size * 1.35)
            fontsize = _fit_fontsize(rect, new_text, size)

            color = u.meta.get("color", 0)
            if isinstance(color, int):
                text_color = (
                    ((color >> 16) & 255) / 255.0,
                    ((color >> 8) & 255) / 255.0,
                    (color & 255) / 255.0,
                )
            else:
                text_color = (0, 0, 0)

            kwargs: dict = {
                "fontsize": fontsize,
                "color": text_color,
                "align": fitz.TEXT_ALIGN_LEFT,
            }
            if fontfile:
                kwargs["fontfile"] = fontfile
            else:
                kwargs["fontname"] = "helv"

            leftover = page.insert_textbox(rect, new_text, **kwargs)
            if leftover < 0:
                for fs in (fontsize - 1, fontsize - 2, max(4.0, fontsize * 0.65)):
                    if fs < 4:
                        break
                    kwargs["fontsize"] = fs
                    if page.insert_textbox(rect, new_text, **kwargs) >= 0:
                        break

    doc.save(str(output_path), garbage=4, deflate=True)
    doc.close()
    return output_path


def write_simple_pdf_from_units(output_path: Path | str, units: list[TextUnit]) -> Path:
    """Fallback: text-only PDF (e.g. DOCX → PDF)."""
    output_path = Path(output_path).with_suffix(".pdf")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fontfile = _pick_font()
    doc = fitz.open()
    page = doc.new_page()
    y = 50.0
    for u in units:
        text = normalize_uzbek_text(u.translated or u.text).strip()
        if not text:
            continue
        if y > 750:
            page = doc.new_page()
            y = 50.0
        kwargs: dict = {"fontsize": 11}
        if fontfile:
            kwargs["fontfile"] = fontfile
        else:
            kwargs["fontname"] = "helv"
        page.insert_text((50, y), text[:500], **kwargs)
        y += 16
    doc.save(str(output_path))
    doc.close()
    return output_path


def resolve_output_format(input_suffix: str, output_format: str) -> str:
    """Return 'docx' or 'pdf'."""
    fmt = (output_format or "auto").lower()
    if fmt == "auto":
        return "docx" if input_suffix == ".docx" else "pdf"
    if fmt in {"docx", "pdf"}:
        return fmt
    raise ValueError("output_format: auto | docx | pdf")


def write_document(
    input_path: Path | str,
    output_path: Path | str,
    units: list[TextUnit],
    *,
    output_format: str = "auto",
) -> Path:
    """
    Write translated document.

    output_format:
      - auto  → same as input (PDF keeps images)
      - docx  → Word (PDF images copied when possible)
      - pdf   → PDF (preferred for scanned layouts / images)
    """
    path = Path(input_path)
    suffix = path.suffix.lower()
    fmt = resolve_output_format(suffix, output_format)
    out = Path(output_path).with_suffix(f".{fmt}")

    if fmt == "docx":
        if suffix == ".docx":
            return write_docx(path, out, units)
        return write_units_as_new_docx(out, units, source_pdf=path if suffix == ".pdf" else None)

    # pdf
    if suffix == ".pdf":
        return write_pdf(path, out, units)
    return write_simple_pdf_from_units(out, units)
