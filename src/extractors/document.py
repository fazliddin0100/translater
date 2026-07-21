"""Extract text units from DOCX/PDF (with position metadata)."""

from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.utils.units import TextUnit


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield ("body", p)
    for ti, table in enumerate(doc.tables):
        yield from _iter_table_paragraphs(f"table{ti}", table)
    for si, section in enumerate(doc.sections):
        for p in section.header.paragraphs:
            yield (f"header{si}", p)
        for ti, table in enumerate(section.header.tables):
            yield from _iter_table_paragraphs(f"header{si}_t{ti}", table)
        for p in section.footer.paragraphs:
            yield (f"footer{si}", p)
        for ti, table in enumerate(section.footer.tables):
            yield from _iter_table_paragraphs(f"footer{si}_t{ti}", table)


def _iter_table_paragraphs(prefix: str, table: Table):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for pi, p in enumerate(cell.paragraphs):
                yield (f"{prefix}_r{ri}c{ci}p{pi}", p)
            for ni, nested in enumerate(cell.tables):
                yield from _iter_table_paragraphs(f"{prefix}_r{ri}c{ci}_n{ni}", nested)


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def extract_docx(path: Path | str) -> list[TextUnit]:
    doc = Document(str(path))
    units: list[TextUnit] = []
    for idx, (where, paragraph) in enumerate(_iter_paragraphs(doc)):
        text = _paragraph_text(paragraph)
        if not text.strip():
            continue
        units.append(
            TextUnit(
                uid=f"docx-{idx}",
                text=text,
                kind="docx_para",
                meta={"where": where, "index": idx},
            )
        )
    return units


def extract_pdf(path: Path | str) -> list[TextUnit]:
    doc = fitz.open(str(path))
    units: list[TextUnit] = []
    try:
        n = 0
        for page_index, page in enumerate(doc):
            data = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text") or ""
                        if not text.strip():
                            continue
                        bbox = list(span["bbox"])
                        units.append(
                            TextUnit(
                                uid=f"pdf-{n}",
                                text=text,
                                kind="pdf_span",
                                meta={
                                    "page": page_index,
                                    "bbox": bbox,
                                    "size": float(span.get("size") or 11),
                                    "color": span.get("color", 0),
                                },
                            )
                        )
                        n += 1
    finally:
        doc.close()
    return units


def extract_document(path: Path | str) -> list[TextUnit]:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    raise ValueError(f"Qo'llab-quvvatlanmaydigan format: {suffix}")
